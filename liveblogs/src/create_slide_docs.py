#! /usr/bin/env python

# Takes a path to a conference data file to programmatically create each document in
# that will contain links to the slides, speaker notes etc.
# usage: ./src/create_slide_docs.py src/conferences/strange-loop-2019/sessions.json

import os
import re
import sys
import json

from docx import Document
from docx.shared import Inches

if __name__ == '__main__':
    ILLEGAL_PATH_CHARS = re.compile('[]¡;,><#|\'&/!@]*')

    try:
        sessions_file_path = sys.argv[1]
        doc_save_path = os.path.dirname(sessions_file_path) 
    except:
        sys.stderr.write('[error]: Please supply the path to the conference data file\n')
        exit()

    try:
        sessions = json.load(open(sessions_file_path))
    except:
        print('[error]: Conference data file not found at "{}"\n'.format(sessions_file_path))

    session = sessions[0]

    for session in sessions:
        doc_path = '{}/{}.docx'.format(doc_save_path, ILLEGAL_PATH_CHARS.sub('', session['title']))
        document = Document()
        document.add_heading(session['title'], level=1)
        document.add_paragraph('')
        document.save(doc_path)
